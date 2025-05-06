from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.shared import Inches
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table


from PIL import Image
from io import BytesIO

from datetime import datetime
from dataclasses import dataclass
from typing import Dict, Tuple
import pandas as pd
import os
import math



import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.figure import Figure 
from io import BytesIO
from docx.shared import Inches
import dataframe_image as dfi
from pathlib import Path


@dataclass(frozen=True)
class ResultOutcome:
    keyword: str
    color: Tuple[int, int, int]  # RGB
    result_ok: bool

# Define all outcomes of a test
EVAL_OUTCOMES: Dict[str, ResultOutcome] = {
    "PASS":         ResultOutcome("PASS",          RGBColor(  0, 200,   0), True),
    "FAIL":         ResultOutcome("FAIL",          RGBColor(200,   0,   0), False),
    "PY ERROR":     ResultOutcome("PY ERROR",      RGBColor(255, 165,   0), False),
    "EXECUTED":     ResultOutcome("EXECUTED",      RGBColor(100, 100, 255), True),
    "LIMIT MISSING":ResultOutcome("LIMIT MISSING", RGBColor(128,   0, 128), False),
    "EVAL NONE":    ResultOutcome("EVAL NONE",     RGBColor(169, 169, 169), False),
    "TYPE ERROR":   ResultOutcome("TYPE ERROR",    RGBColor(255,  99,  71), False),
    "BOOL ERROR":   ResultOutcome("BOOL ERROR",    RGBColor(255, 140,   0), False),
    "FLOAT ERROR":  ResultOutcome("FLOAT ERROR",   RGBColor( 70, 130, 180), False),
    "NO EVAL":      ResultOutcome("NO EVAL",       RGBColor(105, 105, 105), False),
    "NOT EXECUTED": ResultOutcome("NOT EXECUTED",  RGBColor(105, 105, 105), False),
}


@dataclass(frozen=True)
class TC_RUN_SETTING:
    keyword: str
    color: Tuple[int, int, int]
    enabled: bool

# Define all run settings
RUN_SETTINGS: Dict[str, TC_RUN_SETTING] = {
    "ENABLED":          TC_RUN_SETTING("ENABLED",         RGBColor(0, 200, 0),     True),
    "DISABLED":         TC_RUN_SETTING("DISABLED",        RGBColor(169, 169, 169), False),
    "ABORT AFTER FAIL": TC_RUN_SETTING("ABORT",           RGBColor(0, 200, 0),     True),
    "ALWAYS":           TC_RUN_SETTING("ALWAYS",          RGBColor(0, 200, 0),     True),
    "UNKNOWN":          TC_RUN_SETTING("UNKNOWN",         RGBColor(200, 0, 0),     False),
}


@dataclass(frozen=True)
class TC_RUN_STATUS:
    keyword: str
    color: Tuple[int, int, int]
    result_ok: bool

# Define all run statuses
RUN_STATUSES: Dict[str, TC_RUN_STATUS] = {
    "DISABLED":     TC_RUN_STATUS("DISABLED",     RGBColor(169, 169, 169), True),
    "NOT EXECUTED": TC_RUN_STATUS("SKIPPED",      RGBColor(255, 140, 0),   False),
    "PASS":         TC_RUN_STATUS("PASS",         RGBColor(0, 200, 0),     True),
    "FAIL":         TC_RUN_STATUS("FAIL",         RGBColor(200, 0, 0),     False),
    "ERROR":        TC_RUN_STATUS("ERROR",        RGBColor(255, 140, 0),   False),
}

class Cell:
    def __init__(self, table, row, col):
        self.cell = table.rows[row].cells[col]

    def getCell(self):
        return self.cell.text

    def setCell(self, value=None, bold=False, color=None, alignment=None):
        if (value == None):
            value = "/"
            
        try:
            # Access the cell and set its text
            if isinstance(value, float) or isinstance(value, int):
                value = self._format_number(value)
            self.cell.text = str(value).rstrip('\r\n')
    
            # Apply formatting to the cell's text
            for paragraph in self.cell.paragraphs:
                if alignment:
                    paragraph.alignment = alignment
                for run in paragraph.runs:
                    run.bold = bold
                    if color:
                        run.font.color.rgb = color
        except Exception as e:
            print(f"Error in function set Cell {e}")


    def _format_number(self, num):
        """
        Formats a number to have:
        - 3 digits after the decimal point.
        - Uses scientific notation for numbers smaller than 0.001 or larger than 999999.
        - Excludes the exponent if the power is zero.
        """
        if num == 0:
            return "0.000"  # Special case for zero
    
        abs_num = abs(num)
        
        # If the number is within the normal range, return with 3 decimal places
        if 0.001 <= abs_num <= 999999:
            return f"{num:.3f}"
    
        # For numbers outside the range, format in scientific notation with 3 decimal places
        exponent = int(math.floor(math.log10(abs_num)))
        mantissa = num / (10 ** exponent)
    
        # Add a '+' sign for positive exponents
        exponent_str = f"+{exponent}" if exponent > 0 else f"{exponent}"
        return f"{mantissa:.3f}e{exponent_str}"





class TestStep:
    def __init__(self, Auto_Action_cell, LL_cell, UL_cell, Actual_cell, Pass_Fail_cell,KVP_cell):
        self.Auto_Action_cell = Auto_Action_cell
        self.LL_cell = LL_cell
        self.UL_cell = UL_cell
        self.Actual_cell = Actual_cell
        self.Pass_Fail_cell = Pass_Fail_cell
        self.KVP_cell = KVP_cell
        self.testStatus = EVAL_OUTCOMES["NOT EXECUTED"]

    def execute(self, context):
        """
        Executes the command using the given context.
        Stores the result in self.result.
        This method should be customized based on what 'Command' and 'context' mean.
        """
        try:
            self.result = self.Command(context)
        except Exception as e:
            self.result = f"Error during execution: {e}"

    def getResult(self):
        return self.result

    def print(self):
        auto_action = self.Auto_Action_cell.getCell()[:100]  # truncate to 100 chars
        print(
            f"{'KVP:'}{self.KVP_cell.getCell():<10}\t"
            f"{auto_action:<100}\t"
            f"{'LL:'}{self.LL_cell.getCell():<10}\t"
            f"{'UL:'}{self.UL_cell.getCell():<10}\t"
            f"{'Actual:'}{self.Actual_cell.getCell():<10}\t"
            f"{'Pass/Fail:'}{self.Pass_Fail_cell.getCell():<10}"
        )

    def runTestStep(self,context):
        self.result, self.stdout , exec_successful = context.execute(self.Auto_Action_cell.getCell())

        ####### step 1: store result in the word document if the return format is recognized
        if isinstance(self.result, (str,bool, int, float)):
            self.Actual_cell.setCell(self.result, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        elif isinstance(self.result, Image.Image):           
            # Use Auto_Action_cell intentionally due to layout
            paragraph = self.Auto_Action_cell.cell.paragraphs[0]
            run = paragraph.add_run()
    
            buf = BytesIO()
    
            # Ensure Word-compatible format
            image = self.result.convert("RGB") if self.result.mode != "RGB" else self.result
            image.save(buf, format='PNG')
            buf.seek(0)
    
            run.add_picture(buf, width=Inches(4))
        elif isinstance(self.result, Figure):
            # Convert the figure to a PNG image buffer
            buf = BytesIO()
            self.result.savefig(buf, format='png', bbox_inches='tight')
            plt.close(self.result)
            buf.seek(0)
        
            # Insert into Word document
            paragraph = self.Auto_Action_cell.cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(buf, width=Inches(4))
            
        elif isinstance(self.result, pd.DataFrame):
            df_to_render = self.result.head(20)
        
            fig, ax = plt.subplots(figsize=(10, 0.5 * len(df_to_render)))  # Dynamic height
            ax.axis('off')
            table = ax.table(cellText=df_to_render.values, colLabels=df_to_render.columns, loc='center')
            table.scale(1, 1.5)
        
            buf = BytesIO()
            plt.savefig(buf, format='png', bbox_inches='tight', dpi=150)
            plt.close(fig)
            buf.seek(0)
        
            paragraph = self.Auto_Action_cell.cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(buf, width=Inches(4))
            
        #### step 2: evaluate the test limits 
        LL = self.LL_cell.getCell()
        UL = self.UL_cell.getCell()
        
        # test execution failed due to python error:
        if exec_successful == False:
            self.testStatus = EVAL_OUTCOMES['PY ERROR']
            

        # test execution was ok, but no limits are present. Do not ecaluate the limits
        elif LL == '' and UL == '':
            self.testStatus = EVAL_OUTCOMES['EXECUTED']
            
        # one of the limits is missing, we need both for evaluation
        elif LL != '' and UL == '' or LL == '' and UL != '' :
            self.testStatus = EVAL_OUTCOMES['LIMIT MISSING']
            
        # limits are defined, but there is no return from function
        elif self.result == None:
            self.testStatus = EVAL_OUTCOMES['EVAL NONE']

        #check if the value 'can' be compared, we support the following types: str, bool, int, float
        elif not isinstance(self.result, (str, bool, int, float)):
            self.testStatus = EVAL_OUTCOMES['TYPE ERROR']

        # evaluate bool test
        elif isinstance(self.result, bool): 
            if LL.lower() in {'true', 'false'} and UL.lower() in {'true', 'false'}:
                LL = LL.lower() == 'true'
                UL = UL.lower() == 'true'
                if LL == self.result or self.result == UL:
                    self.testStatus = EVAL_OUTCOMES['PASS']
                else:
                    self.testStatus = EVAL_OUTCOMES['FAIL']
            else:
                self.testStatus = EVAL_OUTCOMES['BOOL ERROR']
    
        elif isinstance(self.result, str):
            self.result = self.result.rstrip('\r\n')
            if self.result.lower() in {LL.lower(), UL.lower()}:
                self.testStatus = EVAL_OUTCOMES['PASS']
            else:
                self.testStatus = EVAL_OUTCOMES['FAIL']
    
        elif isinstance(self.result, float) or isinstance(self.result, int):
            try:
                LL_float = float(LL)
                UL_float = float(UL)

                if LL_float <= self.result <= UL_float or UL_float <= self.result <= LL_float:
                    self.testStatus = EVAL_OUTCOMES['PASS']
                else:
                    self.testStatus = EVAL_OUTCOMES['FAIL']

            except ValueError:
                self.testStatus = EVAL_OUTCOMES['FLOAT ERROR']
        else:
            self.testStatus = EVAL_OUTCOMES['NO EVAL']


        self.Pass_Fail_cell.setCell(value=self.testStatus.keyword, bold=False, color=self.testStatus.color, alignment=WD_ALIGN_PARAGRAPH.CENTER)  


class TestCase:
    def __init__(self,table):
        self.table = table     # table of the word document
        self.tests = []        # array to collect all valid test cases

        ############## step1: evaluate to check if all desired tags are present in table. If so, we concider this as a test table (required) ##############
        tags_TestCase = ["<Title>","<TEST_DATE>", "<GROUP>","<MODE>"]
        tags_TestStep = ["<KVP>", "<PYTHON>", "<LL>","<Actual>", "<UL>", "<Pass/fail>"]
        tag_positions = {}
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for tag in tags_TestCase + tags_TestStep:
                    if tag in cell.text:
                        tag_positions[tag] = {"row": row_idx, "col": col_idx}
    
        # Check if all tags are found and display missing tags if there are any
        missing_tags = set(tags_TestCase + tags_TestStep) - set(tag_positions.keys())
        if missing_tags:
            print("Not all tags are detected in this table, skipping table")
            print(f"Missing tags: {', '.join(missing_tags)}")
            self.ValidTC = False
        else:
        ############## step2 :evaluation if the test case tags are in the same row (required) ##############
            rows_with_tags_TestStep = {tag_positions[tag]["row"] for tag in tags_TestStep if tag in tag_positions}
            if len(rows_with_tags_TestStep) != 1:
                print("Test case tags must be on the same row, skipping table")
                self.ValidTC = False
            else:
                self.ValidTC = True
        ############## step3 :map TC cells ##############
                self.cell_title     = Cell(self.table, tag_positions["<Title>"]    ["row"], tag_positions["<Title>"]    ["col"]+1)
                self.cell_group     = Cell(self.table, tag_positions["<GROUP>"]    ["row"], tag_positions["<GROUP>"]    ["col"]+1)
                self.cell_test_date = Cell(self.table, tag_positions["<TEST_DATE>"]["row"], tag_positions["<TEST_DATE>"]["col"]+1)
                self.cell_mode       = Cell(self.table, tag_positions["<MODE>"]    ["row"], tag_positions["<MODE>"]     ["col"]+1)

        ############## step4 :import all tests ##############        
                start_row       = tag_positions["<KVP>"]["row"] + 1  # Start from the row below "<Automated action>"
                kvp_col         = tag_positions["<KVP>"]["col"]
                python_col      = tag_positions["<PYTHON>"]["col"]
                ll_col          = tag_positions["<LL>"]["col"]
                ul_col          = tag_positions["<UL>"]["col"]
                pass_fail_col   = tag_positions["<Pass/fail>"]["col"]
                actual_col      = tag_positions["<Actual>"]["col"]
                
                # Iterate over each row below the "<Automated action>" row in the same column
                for row in range(start_row, len(self.table.rows)):
                    Auto_Action_cell = Cell(self.table, row, python_col)
                    if Auto_Action_cell.getCell() == '':
                        continue
                    LL_cell          = Cell(self.table, row, ll_col)
                    UL_cell          = Cell(self.table, row, ul_col)
                    Actual_cell      = Cell(self.table, row, actual_col)
                    Pass_Fail_cell   = Cell(self.table, row, pass_fail_col)
                    KVP_cell         = Cell(self.table, row, kvp_col)
                    TC = TestStep(Auto_Action_cell, LL_cell, UL_cell, Actual_cell, Pass_Fail_cell,KVP_cell)
                    self.tests.append(TC)


        ############## step5 :set the runSetting and runStatus of the test case ############## 
                mode = self.cell_mode.getCell().strip().lower()
                if mode in {'always','must','run always','runalways'}: 
                    self.runSetting = RUN_SETTINGS["ALWAYS"]
                    self.runStatus = RUN_STATUSES["NOT EXECUTED"]
                elif mode in {'yes', 'true', 'go', 'ok', 'run'}:
                    self.runSetting = RUN_SETTINGS["ENABLED"]
                    self.runStatus = RUN_STATUSES["NOT EXECUTED"]
                elif mode in {'skip after fail','abort','abort after fail','stop after fail','halt after fail','skip on fail','abort on fail','stop on fail','halt on fail'}:
                    self.runSetting = RUN_SETTINGS["ABORT AFTER FAIL"]
                    self.runStatus = RUN_STATUSES["NOT EXECUTED"]
                elif mode in {'no', 'false', 'skip', 'nok', 'disabled'}:
                    self.runSetting = RUN_SETTINGS["DISABLED"]
                    self.runStatus = RUN_STATUSES["DISABLED"]
                else:
                    self.runSetting = RUN_SETTINGS["UNKNOWN"]
                    self.runStatus = RUN_STATUSES["ERROR"]
    
    def print(self):
        print(self.cell_title.getCell(), self.runSetting.keyword, self.runStatus.keyword)
        for test in self.tests:
            test.print()
            
    def runTestCase(self,context):
        errorDetected = False
        failDetected = False
        
        
        if self.runStatus == RUN_STATUSES["NOT EXECUTED"]:
            self.cell_test_date.setCell(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            for test in self.tests:

                # execute the test in the python context
                test.runTestStep(context)

                # see if the test was evaluated as a fail or an execution error
                if test.testStatus.keyword == "FAIL":
                    failDetected = True
                elif test.testStatus.result_ok == False:
                    errorDetected = True

            if errorDetected == True:
                self.runStatus = RUN_STATUSES["ERROR"]
            elif failDetected == True:
                self.runStatus = RUN_STATUSES["FAIL"]
            else:
                self.runStatus = RUN_STATUSES["PASS"]
                

    def getScore(self):
        pass_count = 0
        
        for test in self.tests:
            if test.testStatus.result_ok == True:
                pass_count += 1

        result_summary = f'{pass_count}/{len(self.tests)} PASSED'
        return (pass_count , len(self.tests) , result_summary)         



class testInstance:
    def __init__(self,docPath,context):
        self.docPath = docPath
        self.doc = Document(self.docPath)
        self.TCtables = []

        self.context = context
        

        for tabel in self.doc.tables:
            TC = TestCase(tabel)
            if TC.ValidTC == True:
                self.TCtables.append(TC)

    def addTestSummaryFrontPage(self):
        doc = self.doc
    
        # Create new document body items in memory
        elements_to_prepend = []
    
        
        # Extract full file path and serial number
        full_path = str(Path(self.docPath).resolve())
        serial = self.context.DUT_SerialNumber
        
        # Title paragraph (centered)
        header_paragraph = doc.add_paragraph()
        header_run = header_paragraph.add_run("Test Summary")
        header_run.bold = True
        header_run.font.size = Pt(16)
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elements_to_prepend.append(header_paragraph)
        
        # Info paragraph (left-aligned)
        info_paragraph = doc.add_paragraph()
        info_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        serial_run = info_paragraph.add_run(f"Serial Number: {serial}\n")
        file_run = info_paragraph.add_run(f"File: {full_path}")
        
        # Smaller font for info
        serial_run.font.size = Pt(10)
        file_run.font.size = Pt(10)
        
        elements_to_prepend.append(info_paragraph)


    
        # Define columns and widths
        columns = ['group', 'title', 'Result', 'Summary', 'test_date', 'mode']
        column_widths = [Inches(0.6), Inches(2.5), Inches(0.8), Inches(1.5), Inches(2.0), Inches(1.0)]
    
        # Create table
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
        # Header row
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(columns):
            paragraph = hdr_cells[i].paragraphs[0]
            run = paragraph.add_run(col_name)
            run.bold = True
    
        # Data rows
        for TC in self.TCtables:
            row_cells = table.add_row().cells
            values = [
                TC.cell_group.getCell(),
                TC.cell_title.getCell(),
                TC.runStatus.keyword,
                TC.getScore()[2],
                TC.cell_test_date.getCell(),
                TC.runSetting.keyword,
            ]
            colors = {
                'Result': TC.runStatus.color,
                'mode': TC.runSetting.color,
            }
    
            for i, (value, column_name) in enumerate(zip(values, columns)):
                cell = row_cells[i]
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(str(value))
    
                # Apply RGB color formatting based on dataclass info
                if column_name in colors:
                    r, g, b = colors[column_name]
                    run.font.color.rgb = RGBColor(r, g, b)
    
        # Set column widths
        for row in table.rows:
            for idx, width in enumerate(column_widths):
                row.cells[idx].width = width
    
        elements_to_prepend.append(table)
    
        # Page break
        page_break_para = doc.add_paragraph()
        page_break_para.add_run().add_break(WD_BREAK.PAGE)
        elements_to_prepend.append(page_break_para)
    
        # Insert all at top of body
        body = doc._body._element
        for el in reversed(elements_to_prepend):
            body.insert(0, el._element)


    def run(self, selectedGroups = None):
        a_abort_test_has_failed = False

        
        for TC in self.TCtables:            
            if selectedGroups == None or TC.cell_group.getCell() in selectedGroups.split(','):        # check if the test is in a selected group
                if TC.runSetting == RUN_SETTINGS["ALWAYS"]: 
                    TC.runTestCase(self.context)
                elif TC.runSetting == RUN_SETTINGS["ENABLED"] and a_abort_test_has_failed == False: 
                    TC.runTestCase(self.context)
                elif TC.runSetting == RUN_SETTINGS["ABORT AFTER FAIL"] and a_abort_test_has_failed == False: 
                    TC.runTestCase(self.context)
                    if TC.runStatus != RUN_STATUSES["PASS"]:
                        
                        a_abort_test_has_failed = True
                elif TC.runSetting == RUN_SETTINGS["DISABLED"]:
                    print('test is disabled')
            else:
                TC.runSetting = RUN_SETTINGS["DISABLED"]
                TC.runStatus  = RUN_STATUSES["DISABLED"]
                

    def getTestSummary(self):
        # Collect all entries in a list first
        entries = []
        
        for TC in self.TCtables:
            entry = {
                'title': TC.cell_title.getCell(),
                'test_date': TC.cell_test_date.getCell(),
                'group': TC.cell_group.getCell(),
                'mode': TC.runSetting.keyword,
                'Result': TC.runStatus.keyword,
                'Summary': TC.getScore()[2]
                
            }
            entries.append(entry)
        
        # Create the DataFrame from all entries at once
        df = pd.DataFrame(entries)
        return df

    def getTestDetails(self) -> pd.DataFrame:
        data = []
    
        for table in self.TCtables:
            for test in table.tests:
                data.append({
                    'KVP': test.KVP_cell.getCell(),
                    'Action': test.Auto_Action_cell.getCell(),
                    'LL': test.LL_cell.getCell(),
                    'UL': test.UL_cell.getCell(),
                    'Actual': test.Actual_cell.getCell(),
                    'Pass/Fail': test.Pass_Fail_cell.getCell()
                })
        return pd.DataFrame(data)
        
    def getTestKVP(self) -> pd.DataFrame:
        data = []
        for table in self.TCtables:
            for test in table.tests:
                kvp = test.KVP_cell.getCell()
                if kvp:  # Only include if there's content
                    data.append({
                        'KVP': kvp,
                        'Action': test.Auto_Action_cell.getCell(),
                        'LL': test.LL_cell.getCell(),
                        'UL': test.UL_cell.getCell(),
                        'Actual': test.Actual_cell.getCell(),
                        'Pass/Fail': test.Pass_Fail_cell.getCell()
                    })
    
        return pd.DataFrame(data)

    def saveFile(self, path, openFile = False):
        self.doc.save(path)
        if openFile == True:
            os.startfile(path)